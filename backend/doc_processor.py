from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
try:
    import pandas as pd
except Exception:
    pd = None
import tempfile
import re
import os
from datetime import datetime

# 表头字段
COLUMNS = ['序号', '值班助理', '日期', '上书量（本）', '纠错量（本）',
           '整架范围/整架号', '工作地点', '值班签到', '督导检查情况', '督导检查情况2']
ROW_TEXT_FIELDS = ['值班助理', '日期', '整架范围', '工作地点', '值班签到', '督导检查情况']
ROW_NUMBER_FIELDS = ['上书量', '纠错量']
ROW_FIELDS = ROW_TEXT_FIELDS + ROW_NUMBER_FIELDS

def _set_seq_field(cell, seq_name: str = 'DutySeq') -> None:
    """在单元格中插入 Word SEQ 字段，用于自动编号。"""
    cell.text = ''
    p = cell.paragraphs[0]

    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f'SEQ {seq_name} \\* ARABIC')

    r = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:eastAsia'), '宋体')
    r_fonts.set(qn('w:ascii'), '宋体')
    r_fonts.set(qn('w:hAnsi'), '宋体')
    r_fonts.set(qn('w:cs'), '宋体')
    r_pr.append(r_fonts)
    r_sz = OxmlElement('w:sz')
    r_sz.set(qn('w:val'), '24')  # 12pt
    r_pr.append(r_sz)
    r.append(r_pr)

    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)

    fld.append(r)
    p._p.append(fld)

def _cell_text(cells: list, index: int) -> str:
    if index >= len(cells):
        return ''
    return cells[index].strip()

def _parse_int(value: str) -> int:
    if not value:
        return 0
    m = re.search(r'\d+', value)
    return int(m.group()) if m else 0

def _normalize_rows(rows: list) -> list:
    """规范化行数据字段与类型，确保导出/计算稳定。"""
    normalized = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        normalized_row = {}
        for field in ROW_FIELDS:
            value = row.get(field, '')
            if field in ROW_NUMBER_FIELDS:
                try:
                    normalized_row[field] = int(str(value).strip() or 0)
                except Exception:
                    normalized_row[field] = 0
            else:
                normalized_row[field] = str(value).strip() if value is not None else ''
        normalized.append(normalized_row)
    return normalized

def _compute_totals(rows: list) -> dict:
    """从行数据计算汇总字段。"""
    names = [r.get('值班助理', '').strip() for r in rows if r.get('值班助理')]
    return {
        '总人数': len(set(names)),
        '总班次': len(rows),
        '上书量合计': sum(r.get('上书量', 0) or 0 for r in rows),
        '纠错量合计': sum(r.get('纠错量', 0) or 0 for r in rows),
    }

def parse_single_document(file_path: str) -> list:
    """解析单个Word文档，提取表格数据"""
    doc = Document(file_path)
    if not doc.tables:
        return []

    table = doc.tables[0]
    rows_data = []

    for i, row in enumerate(table.rows):
        if i == 0:  # 跳过表头
            continue

        cells = [cell.text.strip() for cell in row.cells]

        c0 = _cell_text(cells, 0)
        c1 = _cell_text(cells, 1)
        c2 = _cell_text(cells, 2)

        # 跳过总计行和备注行
        if c0 in ['总\n计', '总计', '备注'] or (c1 == '人' and c2 == '次'):
            continue

        # 跳过空行
        if not c1:  # 助理姓名为空则跳过
            continue

        rows_data.append({
            '序号': c0,
            '值班助理': c1,
            '日期': c2,
            '上书量': _parse_int(_cell_text(cells, 3)),
            '纠错量': _parse_int(_cell_text(cells, 4)),
            '整架范围': _cell_text(cells, 5),
            '工作地点': _cell_text(cells, 6),
            '值班签到': _cell_text(cells, 7),
            '督导检查情况': _cell_text(cells, 8),
        })

    return rows_data

def extract_problems(text: str) -> str:
    """提取督导检查情况中的问题"""
    match = re.search(r'存在问题[：:](.*)', text, re.DOTALL)
    return match.group(1).strip() if match else ''

def _parse_date_for_sort(value: str) -> datetime:
    """宽容解析日期字符串，用于稳定排序"""
    if not value:
        return datetime.max
    txt = value.strip()
    # 标准格式
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日"):
        try:
            return datetime.strptime(txt, fmt)
        except Exception:
            pass
            
    # 中文格式：如 11月3日 或 11月03日 (需处理 /星期 或其他后缀)
    # 例如：11.06/四 或 11月04/二
    m = re.search(r'(?P<m>\d{1,2})[月.](?P<d>\d{1,2})', txt)
    if m:
        year = datetime.now().year
        try:
            return datetime(year, int(m.group('m')), int(m.group('d')))
        except Exception:
            return datetime.max
            
    # 备用：仅数字形式 mm-dd 或 mm/dd
    for fmt in ("%m-%d", "%m/%d", "%m.%d"):
        try:
            dt = datetime.strptime(txt, fmt)
            return datetime(datetime.now().year, dt.month, dt.day)
        except Exception:
            pass
    return datetime.max

def parse_documents(file_paths: list) -> dict:
    """解析多个文档并整合数据"""
    all_rows = []
    all_problems = []

    for path in file_paths:
        rows = parse_single_document(path)
        all_rows.extend(rows)

        # 提取问题
        for row in rows:
            problem = extract_problems(row.get('督导检查情况', ''))
            if problem:
                all_problems.append(problem.strip())

    if not all_rows:
        return {'rows': [], 'totals': {}, 'problems': ''}

    # 去重：值班助理 + 日期 + 整架范围 + 工作地点 完全一致时视为重复
    # 注意：这里保留第一条出现的记录
    seen_records = set()
    unique_rows = []
    for row in all_rows:
        # 构造唯一键
        key = (
            row.get('值班助理', '').strip(),
            row.get('日期', '').strip(),
            row.get('整架范围', '').strip(),
            row.get('工作地点', '').strip()
        )
        if key not in seen_records:
            seen_records.add(key)
            unique_rows.append(row)

    if pd is None:
        rows_out = _normalize_rows(unique_rows)
        name_key = ROW_TEXT_FIELDS[0] if ROW_TEXT_FIELDS else ''
        date_key = ROW_TEXT_FIELDS[1] if len(ROW_TEXT_FIELDS) > 1 else ''
        try:
            rows_out.sort(
                key=lambda r: (
                    r.get(name_key, ''),
                    _parse_date_for_sort(r.get(date_key, ''))
                )
            )
        except Exception:
            rows_out.sort(key=lambda r: r.get(name_key, ''))

        totals = _compute_totals(rows_out)

        merged_problems = ''
        if all_problems:
            seen = set()
            dedup_list = []
            for p in all_problems:
                key = re.sub(r'\\s+', '', p)
                if key and key not in seen:
                    seen.add(key)
                    dedup_list.append(p)
            merged_problems = '\\n'.join(dedup_list)

        return {
            'rows': rows_out,
            'totals': totals,
            'problems': merged_problems
        }
    
    # 转为DataFrame进行排序与字段清洗
    df = pd.DataFrame(unique_rows)
    for col in ROW_TEXT_FIELDS:
        if col not in df:
            df[col] = ''
        df[col] = df[col].fillna('').astype(str)
    for col in ROW_NUMBER_FIELDS:
        if col not in df:
            df[col] = 0
        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).astype(int)

    # 构造排序辅助列，按姓名 + 日期排序
    try:
        df['__date_sort__'] = df['日期'].apply(_parse_date_for_sort)
    except Exception:
        df['__date_sort__'] = datetime.max
    df = df.sort_values(['值班助理', '__date_sort__'], ascending=[True, True], kind='mergesort')
    df = df.drop(columns=['__date_sort__'])

    rows_out = _normalize_rows(df.to_dict('records'))
    totals = _compute_totals(rows_out)

    # 督导检查情况文本去重并合并
    merged_problems = ''
    if all_problems:
        seen = set()
        dedup_list = []
        for p in all_problems:
            key = re.sub(r'\s+', '', p)
            if key and key not in seen:
                seen.add(key)
                dedup_list.append(p)
        merged_problems = '\n'.join(dedup_list)

    return {
        'rows': rows_out,
        'totals': totals,
        'problems': merged_problems
    }

def _resolve_template_path() -> str:
    """优先使用后端目录中的模板，避免部署时找不到根目录模板。"""
    env_path = os.getenv('TEMPLATE_PATH')
    candidates = []
    if env_path:
        candidates.append(env_path)
    # backend/templates/...
    backend_dir = os.path.dirname(__file__)
    candidates.append(os.path.join(backend_dir, 'templates', '图书管理岗督导工作情况通报(模板).docx'))
    # repo root fallback
    root_dir = os.path.dirname(os.path.dirname(__file__))
    candidates.append(os.path.join(root_dir, '图书管理岗督导工作情况通报(模板).docx'))

    for path in candidates:
        if path and os.path.exists(path):
            return path
    return ''

def export_document(data: dict) -> str:
    """导出汇总文档"""
    # 使用模板文档
    template_path = _resolve_template_path()
    if not template_path:
        raise FileNotFoundError('未找到模板文件：图书管理岗督导工作情况通报(模板).docx')
    doc = Document(template_path)

    rows = _normalize_rows(data.get('rows', []))
    totals = _compute_totals(rows)
    problems = str(data.get('problems') or '').strip()

    if not rows:
        return None

    # 获取模板中的第一张表
    if not doc.tables:
        table = doc.add_table(rows=1, cols=10)
    else:
        table = doc.tables[0]
    # 确保仅保留表头（第0行），后续逐行追加
    while len(table.rows) > 1:
        tr = table._tbl.tr_lst[-1]
        table._tbl.remove(tr)

    # 数据行 - 先添加所有行，记录分组信息
    groups = []  # [(start_index, end_index, name), ...]
    group_start_index = None
    last_name = None

    for i, row in enumerate(rows, start=1):
        table.add_row()
        name = row.get('值班助理', '')

        if name != last_name:
            if group_start_index is not None:
                groups.append((group_start_index, i - 1, last_name))
            group_start_index = i
            last_name = name

        table.rows[i].cells[2].text = row.get('日期', '')
        table.rows[i].cells[3].text = str(row.get('上书量', 0))
        table.rows[i].cells[4].text = str(row.get('纠错量', 0))
        table.rows[i].cells[5].text = row.get('整架范围', '')
        table.rows[i].cells[6].text = row.get('工作地点', '')
        table.rows[i].cells[7].text = row.get('值班签到') or '√'

    # 记录最后一组
    if group_start_index is not None:
        groups.append((group_start_index, len(rows), last_name))

    # 合并督导检查情况列（先合并，再填内容）
    if len(rows) > 0:
        if len(rows) > 1:
            table.rows[1].cells[8].merge(table.rows[len(rows)].cells[8])
            table.rows[1].cells[9].merge(table.rows[len(rows)].cells[9])
        table.rows[1].cells[8].merge(table.rows[1].cells[9])

    # 合并序号和姓名列，并填写内容
    group_index = 1
    for (start_idx, end_idx, name) in groups:
        # 合并多行
        if end_idx > start_idx:
            for col in [0, 1]:
                table.rows[start_idx].cells[col].merge(table.rows[end_idx].cells[col])
        # 填写序号与姓名（一个人一个序号）
        table.rows[start_idx].cells[0].text = str(group_index)
        table.rows[start_idx].cells[1].text = name
        group_index += 1

    # 填写督导检查情况（合并后再填写）
    if len(rows) > 0:
        cell = table.rows[1].cells[8]
        cell.text = ''
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        p1 = cell.paragraphs[0]
        run1 = p1.add_run("一、小组总结")
        run1.bold = True

        cell.add_paragraph("1. 巡视到位")
        cell.add_paragraph("2. 工作认真负责")

        cell.add_paragraph('')
        p4 = cell.add_paragraph()
        run4 = p4.add_run("二、存在问题")
        run4.bold = True

        if problems:
            prob_list = problems.split('\n')
            for idx, prob in enumerate(prob_list, 1):
                if prob.strip():
                    clean_prob = re.sub(r'^(?:[1-9]|[一二三四五六七八九十])[、.]\s*', '', prob.strip())
                    p_prob = cell.add_paragraph(f"{idx}. {clean_prob}")
                    for run in p_prob.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            p_none = cell.add_paragraph("无")
            for run in p_none.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)

    # 总计行
    total_row = table.add_row()
    # 0. 总计
    total_row.cells[0].text = ''
    p_total = total_row.cells[0].paragraphs[0]
    r_total = p_total.add_run('总\n计')
    r_total.bold = True
    
    # 1. 人数 (值班助理列)
    total_row.cells[1].text = f"{totals.get('总人数', 0)}人"
    # 2. 班次 (日期列)
    total_row.cells[2].text = f"{totals.get('总班次', 0)}班"
    # 3. 上书量 (上书量列)
    total_row.cells[3].text = str(totals.get('上书量合计', 0))
    # 4. 纠错量 (纠错量列)
    total_row.cells[4].text = str(totals.get('纠错量合计', 0))
    # 5. 整架范围 (占位)
    total_row.cells[5].text = '-'
    # 6. 工作地点 (占位)
    total_row.cells[6].text = '-'
    # 7. 值班签到 (占位)
    total_row.cells[7].text = '-'
    # 8. 督导检查 (占位)
    total_row.cells[8].text = '-'
    # 9. 督导检查2 (占位)
    total_row.cells[9].text = '-'
    
    # 合并最后两列
    total_row.cells[8].merge(total_row.cells[9])

    # 备注行
    note_row = table.add_row()
    note_row.cells[0].text = ''
    p_note_title = note_row.cells[0].paragraphs[0]
    r_note_title = p_note_title.add_run('备注')
    r_note_title.bold = True
    
    note_text = '每位助理工作情况良好，能很好地兼顾学习和工作，整体的工作状态都不错，但其中仍存在部分不足，希望大家有则改之，无则加勉。'
    note_row.cells[1].text = ''
    p_note_content = note_row.cells[1].paragraphs[0]
    r_note_content = p_note_content.add_run(note_text)
    r_note_content.bold = True
    
    # 合并备注内容单元格
    table.rows[len(rows) + 2].cells[1].merge(table.rows[len(rows) + 2].cells[7])
    note_row.cells[8].text = '值班督导'

    # 设置全局字体、行高与对齐方式
    data_row_start = 1
    data_row_end = len(rows)
    for row_idx, row in enumerate(table.rows):
        is_data_row = data_row_start <= row_idx <= data_row_end
        # 数据行最小高度 1.71cm，允许督导检查情况扩展
        if is_data_row:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Cm(1.71)
        else:
            row.height = Cm(1.71)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for col_idx, cell in enumerate(row.cells):
            is_left_block = col_idx <= 7
            is_supervisor_cell = is_data_row and col_idx >= 8

            if is_left_block:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            elif is_supervisor_cell:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            else:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for paragraph in cell.paragraphs:
                if is_left_block:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif is_supervisor_cell:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:cs'), '宋体')

    # 保存到临时文件
    # 打开文档时自动更新字段（SEQ 编号等）
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name
