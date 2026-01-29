from docx import Document

doc = Document(r"E:\newlife\自动化文档处理系统\测试\督导工作情况汇总——07.docx")
table = doc.tables[0]

print("=== 检查序号列 ===")
for i in range(min(10, len(table.rows))):
    cell0 = table.rows[i].cells[0].text
    cell1 = table.rows[i].cells[1].text
    print(f"行{i}: 序号=[{cell0}] 姓名=[{cell1}]")

print("\n=== 检查督导检查情况列 ===")
cell8 = table.rows[1].cells[8]
print(f"段落数: {len(cell8.paragraphs)}")
for j, p in enumerate(cell8.paragraphs[:10]):
    print(f"段落{j}: [{p.text[:50] if p.text else '(空)'}]")
